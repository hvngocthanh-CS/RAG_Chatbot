"""
Document Parser — Step 1 of the RAG pipeline.

Uses pdfplumber for BOTH text and table extraction from PDFs.
Key capabilities:
  - Header/footer removal via y-coordinate filtering
  - Table regions excluded from text extraction (no duplication)
  - Paragraph-level block granularity (not page-level dumps)
  - Heading detection via font size analysis + regex patterns
"""
import re
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

# PDF page coordinate thresholds (A4: 595 x 842 points)
HEADER_Y_THRESHOLD = 55      # anything above this = header
FOOTER_Y_THRESHOLD = 780     # anything below this = footer


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class TextBlock:
    """A contiguous block of text extracted from a document."""
    text: str
    page_number: Optional[int] = None
    block_type: str = "paragraph"       # "title" | "heading" | "paragraph"
    section: str = ""                   # Nearest heading above this block

    def to_dict(self) -> dict:
        return {
            "text": self.text,
            "page_number": self.page_number,
            "block_type": self.block_type,
            "section": self.section,
        }


@dataclass
class Table:
    """A table extracted from a document."""
    headers: List[str]
    rows: List[List[str]]
    page_number: Optional[int] = None
    name: str = ""
    section: str = ""

    @property
    def is_empty(self) -> bool:
        return len(self.rows) == 0

    def to_dict(self) -> dict:
        return {
            "headers": self.headers,
            "rows": self.rows,
            "page_number": self.page_number,
            "name": self.name,
            "section": self.section,
        }


@dataclass
class ParsedDocument:
    """Complete output of parsing a single file."""
    text_blocks: List[TextBlock] = field(default_factory=list)
    tables: List[Table] = field(default_factory=list)
    page_count: int = 0
    file_type: str = ""

    def to_dict(self) -> dict:
        return {
            "text_blocks": [b.to_dict() for b in self.text_blocks],
            "tables": [t.to_dict() for t in self.tables],
            "metadata": {
                "page_count": self.page_count,
                "file_type": self.file_type,
                "text_block_count": len(self.text_blocks),
                "table_count": len(self.tables),
            },
        }


# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------

_NUMBERED_HEADING_RE = re.compile(r"^\d+(\.\d+)*\.?\s+\S")


def _is_heading(text: str) -> bool:
    """Heuristic: short ALL-CAPS text, or numbered section headers."""
    text = text.strip()
    if not text or len(text) > 150:
        return False
    if text.isupper() and len(text) < 80:
        return True
    # Numbered pattern must also be short — long text starting with "5." is
    # a heading merged with body text, not a real heading.
    if _NUMBERED_HEADING_RE.match(text) and len(text) < 100:
        return True
    # Removed: the old "short text without ending punctuation" rule
    # had too many false positives (bullet items, short phrases,
    # table captions, etc. were all misclassified as headings).
    return False


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

class DocumentParser:
    """
    Multi-format document parser using pdfplumber exclusively for PDFs.

    PDF pipeline per page:
      1. Crop page to remove header/footer by y-coordinate
      2. Find table bounding boxes in cropped area
      3. Extract tables separately via find_tables()
      4. Extract words outside table regions via extract_words()
      5. Group words → lines → paragraphs by vertical gaps
      6. Classify each paragraph as heading or body text via font size + regex
    """

    async def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported format '{ext}'. Supported: {SUPPORTED_EXTENSIONS}"
            )

        logger.info("Parsing %s (%s)", path.name, ext)

        if ext == ".pdf":
            doc = await self._parse_pdf(path)
        elif ext == ".docx":
            doc = await self._parse_docx(path)
        else:
            doc = await self._parse_text(path)

        logger.info(
            "Parsed %s: %d text blocks, %d tables, %d pages",
            path.name, len(doc.text_blocks), len(doc.tables), doc.page_count,
        )
        return doc

    # ------------------------------------------------------------------
    # PDF — pdfplumber only
    # ------------------------------------------------------------------
    async def _parse_pdf(self, path: Path) -> ParsedDocument:
        import pdfplumber

        all_text_blocks: List[TextBlock] = []
        all_tables: List[Table] = []
        current_section = ""

        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):

                # Step 1: Crop — remove header and footer regions
                body = page.crop((
                    0, HEADER_Y_THRESHOLD,
                    page.width, FOOTER_Y_THRESHOLD,
                ))

                # Step 2: Find tables and collect their bounding boxes
                found_tables = body.find_tables()
                table_bboxes = [t.bbox for t in found_tables]

                # Step 3: Extract tables
                for idx, t in enumerate(found_tables):
                    table = self._clean_raw_table(
                        t.extract(), page_num, idx, current_section,
                    )
                    if table and not table.is_empty:
                        all_tables.append(table)

                # Step 4: Extract words outside table regions (with font size)
                words = body.extract_words(
                    keep_blank_chars=False,
                    extra_attrs=["size"],
                )
                body_words = [
                    w for w in words
                    if not self._word_in_any_bbox(w, table_bboxes)
                ]

                if not body_words:
                    continue

                # Step 5: Group words → lines → paragraphs
                body_font_size = self._most_common_size(body_words)
                paragraphs = self._words_to_paragraphs(body_words)

                # Step 6: Classify each paragraph and build TextBlocks
                for para_text, avg_font_size in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue

                    is_heading = (
                        avg_font_size > body_font_size + 1.5
                        or _is_heading(para_text)
                    )
                    if is_heading:
                        current_section = para_text
                        block_type = "heading"
                    else:
                        block_type = "paragraph"

                    all_text_blocks.append(TextBlock(
                        text=para_text,
                        page_number=page_num,
                        block_type=block_type,
                        section=current_section,
                    ))

        return ParsedDocument(
            text_blocks=all_text_blocks,
            tables=all_tables,
            page_count=page_count,
            file_type="pdf",
        )

    # ------------------------------------------------------------------
    # PDF helpers
    # ------------------------------------------------------------------
    @staticmethod
    def _word_in_any_bbox(word: dict, bboxes: list) -> bool:
        """True if the word falls inside any of the given bounding boxes."""
        for (tx0, ttop, tx1, tbottom) in bboxes:
            if (word["x0"] >= tx0 - 2 and word["x1"] <= tx1 + 2
                    and word["top"] >= ttop - 2 and word["bottom"] <= tbottom + 2):
                return True
        return False

    @staticmethod
    def _most_common_size(words: list) -> float:
        """Return the most frequent font size among words (= body text size)."""
        from collections import Counter
        sizes = [round(w.get("size", 10), 1) for w in words]
        if not sizes:
            return 10.0
        return Counter(sizes).most_common(1)[0][0]

    @staticmethod
    def _words_to_paragraphs(words: list) -> List[tuple]:
        """
        Group words into lines (by y-position), then lines into paragraphs
        (by vertical gap). Returns list of (text, avg_font_size).
        """
        if not words:
            return []

        # Sort top-to-bottom, then left-to-right
        words = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))

        # --- Group into lines (words at similar y) ---
        lines = []
        cur_line = [words[0]]
        for w in words[1:]:
            if abs(w["top"] - cur_line[-1]["top"]) < 3:
                cur_line.append(w)
            else:
                lines.append(cur_line)
                cur_line = [w]
        lines.append(cur_line)

        # Build line objects
        line_objs = []
        for lw in lines:
            lw_sorted = sorted(lw, key=lambda w: w["x0"])
            text = " ".join(w["text"] for w in lw_sorted)
            top = min(w["top"] for w in lw_sorted)
            sizes = [w.get("size", 10) for w in lw_sorted]
            avg_size = sum(sizes) / len(sizes)
            line_objs.append({"text": text, "top": top, "avg_size": avg_size})

        # --- Group lines into paragraphs (by vertical gap) ---
        paragraphs = []
        cur_para = [line_objs[0]]

        for i in range(1, len(line_objs)):
            gap = line_objs[i]["top"] - line_objs[i - 1]["top"]
            # Adaptive threshold: 1.8x the expected line height based on
            # the current paragraph's font size.  Falls back to 25pt when
            # font info is unavailable.
            avg_size = cur_para[-1].get("avg_size", 10)
            para_break_threshold = max(avg_size * 1.2 * 1.8, 18)
            if gap > para_break_threshold:
                para_text = " ".join(l["text"] for l in cur_para)
                avg_font = sum(l["avg_size"] for l in cur_para) / len(cur_para)
                paragraphs.append((para_text, avg_font))
                cur_para = [line_objs[i]]
            else:
                cur_para.append(line_objs[i])

        # Last paragraph
        if cur_para:
            para_text = " ".join(l["text"] for l in cur_para)
            avg_font = sum(l["avg_size"] for l in cur_para) / len(cur_para)
            paragraphs.append((para_text, avg_font))

        return paragraphs

    @staticmethod
    def _clean_raw_table(
        raw_table: list, page_number: int, index: int, section: str,
    ) -> Optional[Table]:
        """Convert pdfplumber raw table (list of lists) into a Table dataclass."""
        if not raw_table or len(raw_table) < 2:
            return None

        def clean(cell) -> str:
            if cell is None:
                return ""
            return " ".join(str(cell).split()).strip()

        headers = [clean(c) for c in raw_table[0]]
        rows = []
        for row in raw_table[1:]:
            cleaned = [clean(c) for c in row]
            cleaned = (cleaned + [""] * len(headers))[:len(headers)]
            if any(cleaned):
                rows.append(cleaned)

        if not rows:
            return None

        return Table(
            headers=headers,
            rows=rows,
            page_number=page_number,
            name=f"Table_p{page_number}_{index + 1}",
            section=section,
        )

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------
    async def _parse_docx(self, path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(path))
        text_blocks: List[TextBlock] = []
        tables: List[Table] = []
        current_section = ""

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]

            if tag == "p":
                para = self._find_paragraph(doc, element)
                if not para or not para.text.strip():
                    continue
                style = para.style.name if para.style else ""
                is_heading = "Heading" in style

                if is_heading:
                    current_section = para.text.strip()

                text_blocks.append(TextBlock(
                    text=para.text.strip(),
                    block_type="heading" if is_heading else "paragraph",
                    section=current_section,
                ))

            elif tag == "tbl":
                tbl = self._find_table(doc, element)
                if tbl:
                    table = self._extract_docx_table(tbl, current_section)
                    if table and not table.is_empty:
                        tables.append(table)

        return ParsedDocument(
            text_blocks=text_blocks,
            tables=tables,
            page_count=0,
            file_type="docx",
        )

    @staticmethod
    def _find_paragraph(doc, element):
        for p in doc.paragraphs:
            if p._element is element:
                return p
        return None

    @staticmethod
    def _find_table(doc, element):
        for t in doc.tables:
            if t._element is element:
                return t
        return None

    @staticmethod
    def _extract_docx_table(tbl, section: str) -> Optional[Table]:
        all_rows = []
        for row in tbl.rows:
            all_rows.append([cell.text.strip() for cell in row.cells])
        if len(all_rows) < 2:
            return None
        return Table(
            headers=all_rows[0],
            rows=all_rows[1:],
            name=f"Table_{len(all_rows)}rows",
            section=section,
        )

    # ------------------------------------------------------------------
    # TXT / Markdown
    # ------------------------------------------------------------------
    async def _parse_text(self, path: Path) -> ParsedDocument:
        content = path.read_text(encoding="utf-8", errors="ignore")
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

        text_blocks: List[TextBlock] = []
        current_section = ""
        for para in paragraphs:
            first_line = para.split("\n")[0]
            if first_line.startswith("#"):
                current_section = first_line.lstrip("# ").strip()
                block_type = "heading"
            else:
                block_type = "paragraph"

            text_blocks.append(TextBlock(
                text=para,
                block_type=block_type,
                section=current_section,
            ))

        return ParsedDocument(
            text_blocks=text_blocks,
            tables=[],
            page_count=0,
            file_type=path.suffix.lstrip("."),
        )
